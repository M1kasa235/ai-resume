"""共享文本提取器 — PDF / DOCX / TXT

所有 RAG 入库路径共用此模块，避免 resume 和 knowledge 两套重复代码。
"""

import logging

logger = logging.getLogger(__name__)


def extract_text_from_txt(file_path: str) -> str:
    """提取 TXT 文本"""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read().strip()


def extract_text_from_pdf(file_path: str) -> str:
    """提取 PDF 文本"""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """提取 DOCX 文本，覆盖段落 + 文本框 + 表格三种来源"""
    try:
        from docx import Document
        import xml.etree.ElementTree as ET
    except ImportError:
        raise RuntimeError("需要安装 python-docx: pip install python-docx")

    doc = Document(file_path)
    parts = []

    # 1. 普通段落
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # 2. 文本框（shape / textbox）
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for txbx in doc.element.body.findall('.//w:txbxContent', ns):
        lines = []
        for p in txbx.findall('.//w:p', ns):
            line = ''.join(
                t.text or '' for t in p.findall('.//w:t', ns)
            ).strip()
            if line:
                lines.append(line)
        if lines:
            parts.append('\n'.join(lines))

    # 3. 表格
    for table in doc.tables:
        for row in table.rows:
            row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_parts:
                parts.append('\t'.join(row_parts))

    return '\n\n'.join(parts).strip()


def extract_text_from_file(file_path: str) -> str:
    """自动识别文件类型并提取文本"""
    import os
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        return extract_text_from_txt(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".doc", ".docx"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"不支持的文本文件格式: {ext}，支持 .txt .pdf .doc .docx")